import hashlib
import json
import os
import re
import shutil
import uuid
from collections import defaultdict
from datetime import datetime, timezone
from typing import List, Tuple

import fitz  # PyMuPDF for PDF image extraction
import nest_asyncio
import pandas as pd
import streamlit as st
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from dotenv import load_dotenv
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import (
    Docx2txtLoader,
    PyPDFLoader,
    TextLoader,
    UnstructuredWordDocumentLoader,
)
from langchain_community.vectorstores import Chroma
from langchain_google_genai import GoogleGenerativeAIEmbeddings

nest_asyncio.apply()
load_dotenv()

# --- Constants ---
DEFAULT_DOCS_DIR    = "docs"
DEFAULT_PERSIST_DIR = "Vector_DB - Documents"
DEFAULT_CHUNKS_DIR  = "chunks"
CHUNK_SIZE          = int(os.getenv("CHUNK_SIZE", 2000))
CHUNK_OVERLAP       = int(os.getenv("CHUNK_OVERLAP", 200))


def _parse_cache_line(line: str) -> Tuple[str, List[str]]:
    """Return filename and list of vector IDs from a cache line."""
    raw = line.strip()
    if not raw:
        return "", []
    if "\\" in raw:
        fname, ids_part = raw.split("\\", 1)
        ids = [i for i in ids_part.split("/") if i]
        return fname, ids
    return raw, []


def _format_cache_line(filename: str, ids: List[str]) -> str:
    """Format cache line as filename\\id1/id2/...; falls back to filename only."""
    return f"{filename}\\{'/'.join(ids)}" if ids else filename


def is_gibberish(text, threshold=0.3):
    if not text:
        return True
    alnum = sum(c.isalnum() for c in text)
    ratio = alnum / max(len(text), 1)
    return ratio < threshold


@st.cache_resource
def get_ocr(lang: str):
    from paddleocr import PaddleOCR
    return PaddleOCR(lang=lang, use_angle_cls=True, show_log=False)


def ocr_pdf_with_paddleocr(pdf_path, lang='vi'):  # Vietnamese support
    ocr = get_ocr(lang)
    doc = fitz.open(pdf_path)
    all_text = []
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        images = page.get_images(full=True)
        page_text = []
        for img_index, img in enumerate(images):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            import cv2
            import numpy as np
            img_array = np.frombuffer(image_bytes, np.uint8)
            img_cv = cv2.imdecode(img_array, cv2.IMREAD_COLOR)
            if img_cv is not None:
                result = ocr.ocr(img_cv, cls=True)
                for line in result:
                    for box in line:
                        text = box[1][0]
                        page_text.append(text)
        # If no images, try to render the page as an image and OCR it
        if not images:
            pix = page.get_pixmap()
            img_cv = np.frombuffer(pix.samples, dtype=np.uint8).reshape((pix.height, pix.width, pix.n))
            result = ocr.ocr(img_cv, cls=True)
            for line in result:
                for box in line:
                    text = box[1][0]
                    page_text.append(text)
        if page_text:
            all_text.append(f"Page {page_num+1}:\n" + "\n".join(page_text))
    doc.close()
    if all_text:
        return [Document(
            page_content="\n\n".join(all_text),
            metadata={
                "source": pdf_path,
                "filename": os.path.basename(pdf_path)
            }
        )]
    return []


def load_text_from_txt_file(filepath: str) -> List[Document]:
    encodings = ['utf-8', 'utf-16', 'cp1252', 'iso-8859-1', 'gbk']
    for encoding in encodings:
        try:
            with open(filepath, 'r', encoding=encoding) as f:
                content = f.read()
                return [Document(
                    page_content=content,
                    metadata={
                        "source": filepath,
                        "filename": os.path.basename(filepath),
                        "img_list": ""
                    }
                )]
        except (UnicodeDecodeError, LookupError):
            continue

    st.error(f"Cannot decode text file: {filepath}")
    return []


def load_text_from_docx_file(filepath: str) -> List[Document]:
    dirname = os.path.dirname(filepath)
    basename = os.path.basename(filepath)
    img_dir = os.path.join(dirname, "images", basename)
    os.makedirs(img_dir, exist_ok=True)

    doc = DocxDocument(filepath)
    rels = doc.part.rels
    text_list = []
    img_counter = 0
    img_paths = {}
    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "v": "urn:schemas-microsoft-com:vml",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }

    def _save_image(_rid: str):
        nonlocal img_counter
        # if rid not in rels:
        #     return
        image_part = rels[_rid].target_part
        image_bytes = image_part.blob
        img_counter += 1
        img_name = f"image_{img_counter}.png"
        img_path = os.path.join(img_dir, img_name)
        with open(img_path, "wb") as f:
            f.write(image_bytes)
        text_list.append(f"[IMAGE:{img_name}]")
        img_paths[img_name] = img_path

    for para in doc.paragraphs:
        para_text = ""
        for run in para.runs:
            # normal text
            if run.text:
                para_text += run.text

            # DrawingML images (native DOCX)
            drawings = run._element.findall(".//w:drawing", namespaces=ns)
            if drawings:
                # # flush accumulated text first
                # text_list.append(para_text)
                # para_text = ""
                para_text = para_text.strip()
                if para_text:
                    text_list.append(para_text)
                    para_text = ""
                # extract embedded image
                for drawing in drawings:
                    blips = drawing.findall(".//a:blip", namespaces=ns)
                    for blip in blips:
                        rid = blip.get(qn("r:embed"))
                        if rid:
                            _save_image(rid)

            # VML images (DOC → Spire → DOCX)
            picts = run._element.findall(".//w:pict", namespaces=ns)
            if picts:
                para_text = para_text.strip()
                if para_text:
                    text_list.append(para_text)
                    para_text = ""
                for pict in picts:
                    imagedata = pict.findall(".//v:imagedata", namespaces=ns)
                    for img in imagedata:
                        rid = img.get(qn("r:id"))
                        if rid:
                            _save_image(rid)

        # # flush remaining paragraph text
        # text_list.append(para_text)
        para_text = para_text.strip()
        if para_text:
            text_list.append(para_text)
    
    full_text = "\n\n".join(text_list)
    spire_watermark = "Evaluation Warning: The document was created with Spire.Doc for Python."
    full_text = full_text.replace(spire_watermark, "").strip()

    return [Document(
        page_content=full_text,
        metadata={
            "source": filepath,
            "filename": basename,
            "img_paths_json": json.dumps(img_paths),
            "img_list": ", ".join(img_paths.keys())
        }
    )]


def extract_text(file_list: List[str], docs_dir: str = DEFAULT_DOCS_DIR):
    docs = []
    for fn in file_list:
        path = os.path.join(docs_dir, fn)
        try:
            if fn.lower().endswith(".pdf"):
                loaded = PyPDFLoader(path).load()
                for d in loaded:
                    d.metadata["filename"] = os.path.basename(path)
                    d.metadata.setdefault("img_list", "")
                all_text = " ".join(doc.page_content for doc in loaded)
                if not loaded or is_gibberish(all_text):
                    st.warning(f"⚠️ Falling back to PaddleOCR for: {fn}")
                    try:
                        ocr_loaded = ocr_pdf_with_paddleocr(path, lang='vi')
                        docs.extend(ocr_loaded)
                    except Exception as ocr_e:
                        continue
                else:
                    docs.extend(loaded)
            elif fn.lower().endswith(".txt"):
                docs.extend(load_text_from_txt_file(path))
            elif fn.lower().endswith(".docx"):
                # loaded = Docx2txtLoader(path).load()
                # for d in loaded:
                #     d.metadata["filename"] = os.path.basename(path)
                # docs.extend(loaded)
                docs.extend(load_text_from_docx_file(path))
            # If the code runs as expected, it will never reach this branch
            # because .doc files are already converted to .docx during upload.
            elif fn.lower().endswith(".doc"):
                # loaded = UnstructuredWordDocumentLoader(path).load()
                # for d in loaded:
                #     d.metadata["filename"] = os.path.basename(path)
                #     d.metadata.setdefault("img_list", "")
                # docs.extend(loaded)
                #
                # path = convert_doc2docx(path)
                # docs.extend(load_text_from_docx_file(path))
                with open('tmp_log.txt', 'a', encoding='utf-8') as logf:
                    logf.write(f"Skipping .doc file (not supported): {fn}\n")
            elif fn.lower().endswith(".xls") or fn.lower().endswith(".xlsx"):
                # Excel support for both .xls and .xlsx, with engine selection
                try:
                    if fn.lower().endswith(".xls"):
                        df = pd.read_excel(path, sheet_name=None, engine="xlrd")
                    else:
                        df = pd.read_excel(path, sheet_name=None, engine="openpyxl")
                except Exception as e:
                    # Fallback to default engine if specified engine fails
                    try:
                        df = pd.read_excel(path, sheet_name=None)
                    except Exception as e2:
                        st.error(f"❌ Failed to read Excel file {fn}: {e2}")
                        continue
                text = ""
                for sheet, data in df.items():
                    text += f"Sheet: {sheet}\n"
                    text += data.to_string(index=False)
                    text += "\n\n"
                if text.strip():
                    docs.append(Document(
                        page_content=text,
                        metadata={
                            "source": path,
                            "filename": os.path.basename(path),
                            "img_list": ""
                        }
                    ))
            else:
                st.warning(f"⚠️ Unsupported file type: {fn}")
        except Exception as e:
            st.error(f"❌ Failed to process {fn}: {e}")

    # Ensure all documents carry required metadata keys for downstream prompts
    for d in docs:
        meta = d.metadata or {}
        meta.setdefault("img_list", "")
        meta.setdefault("added_at", datetime.now(tz=timezone.utc).isoformat())
        if "filename" not in meta and meta.get("source"):
            meta["filename"] = os.path.basename(meta["source"])
        d.metadata = meta
    return docs


def get_text_chunks(
    docs,
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP
):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", " ", ""],
    )
    return splitter.split_documents(docs)


def save_text_chunks(
    chunks,
    chunks_dir: str = DEFAULT_CHUNKS_DIR,
    overwrite: bool = True
) -> None:
    if overwrite and os.path.isdir(chunks_dir):
        shutil.rmtree(chunks_dir)
    os.makedirs(chunks_dir, exist_ok=True)

    for i, chunk in enumerate(chunks):
        src = chunk.metadata.get("source", "")
        base = os.path.splitext(os.path.basename(src))[0] if src else "doc"
        fname = f"{base}_chunk_{i:04d}.txt"
        path = os.path.join(chunks_dir, fname)
        with open(path, "w", encoding="utf-8") as f:
            f.write(chunk.page_content)

    print(f"✅ Exported {len(chunks)} chunks to '{chunks_dir}/'")


def hash_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


# --- User-specific Constants ---
def get_user_dirs(username: str):
    """Get user-specific directory paths"""
    user_base = f"data/kb/{username}"
    return {
        'docs': f"{user_base}/docs",
        'chunks': f"{user_base}/chunks",
        'vectordb': f"{user_base}/vector_db"
    }


def ensure_user_dirs(username: str):
    """Create user-specific directories"""
    dirs = get_user_dirs(username)
    for dir_path in dirs.values():
        os.makedirs(dir_path, exist_ok=True)
    return dirs


def has_new_files_user(username: str, current_files: List[str]) -> bool:
    """Check for new files in user's directory"""
    dirs = get_user_dirs(username)
    cache_path = os.path.join(dirs['vectordb'], "files.txt")

    if not os.path.exists(cache_path):
        return True

    with open(cache_path, "r", encoding="utf-8") as f:
        cached_files = set(_parse_cache_line(line)[0] for line in f.readlines())
    return set(current_files) != cached_files


def get_vectorstore_user(
        username: str,
        file_list: List[str] = []
) -> Chroma:
    """
    Get user-specific vectorstore
    """

    # Ensure user directories exist
    dirs = ensure_user_dirs(username)

    embedding = GoogleGenerativeAIEmbeddings(
        model=os.getenv('TEXT_EMBEDDING_MODEL'),
        google_api_key=os.getenv('GOOGLE_API_KEY')
    )

    # Load or create user-specific vectorstore
    vectordb = Chroma(
        persist_directory=dirs['vectordb'],
        embedding_function=embedding
    )

    # Load previously embedded file list
    cache_path = os.path.join(dirs['vectordb'], "files.txt")
    prev_files = set()
    if os.path.exists(cache_path):
        with open(cache_path, "r", encoding="utf-8") as f:
            prev_files = set(_parse_cache_line(line)[0] for line in f)

    # Filter new files
    new_files = [f for f in file_list if f not in prev_files]
    if not new_files:
        return vectordb

    st.info(f"🆕 Processing {len(new_files)} new files for user: {username}")

    # Extract text from new files
    docs = extract_text(new_files, dirs['docs'])
    chunks = get_text_chunks(docs)

    # Deduplicate by chunk content hash
    seen_hashes = set()
    unique_chunks: List[Document] = []
    ids_by_file: defaultdict[str, List[str]] = defaultdict(list)
    for chunk in chunks:
        content_hash = hash_text(chunk.page_content)
        if content_hash not in seen_hashes:
            seen_hashes.add(content_hash)
            unique_chunks.append(chunk)

    # Add only unique chunks
    if unique_chunks:
        all_ids: List[str] = []

        # Assign UUIDs per chunk and keep track by source filename
        for chunk in unique_chunks:
            src_path = chunk.metadata.get("source", "")
            fname = os.path.basename(src_path)
            cid = str(uuid.uuid4())
            ids_by_file[fname].append(cid)
            all_ids.append(cid)

        vectordb.add_documents(unique_chunks, ids=all_ids)
        vectordb.persist()

        # st.success(f"✅ Added {len(unique_chunks)} unique chunks for {username}")

        # Lưu thông báo vào session state thay vì st.success
        st.session_state[f'vectorstore_success_{username}'] = f"✅ Added {len(unique_chunks)} unique chunks for {username}"

    # Update file cache
    with open(cache_path, "a", encoding="utf-8") as f:
        for fname in new_files:
            ids = ids_by_file.get(fname, [])
            f.write(_format_cache_line(fname, ids) + "\n")

    # Save chunks for inspection
    save_text_chunks(unique_chunks, chunks_dir=dirs['chunks'], overwrite=False)

    return vectordb


def cleanup_user_data(username: str):
    """Clean up all user data"""
    user_base = f"data/kb/{username}"
    if os.path.exists(user_base):
        shutil.rmtree(user_base)
        st.success(f"🗑️ Cleaned up all data for user: {username}")


# def rebuild_user_vectorstore(username: str):
#     """Rebuild vectorstore from scratch for user"""
#     dirs = get_user_dirs(username)
#
#     # Xóa toàn bộ vectorstore cũ
#     if os.path.exists(dirs['vectordb']):
#         shutil.rmtree(dirs['vectordb'])
#
#     # Tạo lại thư mục
#     os.makedirs(dirs['vectordb'], exist_ok=True)
#
#     # Lấy danh sách file hiện tại và rebuild
#     current_files = get_user_documents(username)
#     if current_files:
#         return get_vectorstore_user(username, current_files)
#
#     return None
